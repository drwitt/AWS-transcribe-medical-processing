"""
Author: Danny Witt
Purpose: read s3 bucket with .json files, apply json parser, and 
    deposit resulting word (.docx) output file in output s3 bucket
Input: s3 bucket path with transcribed json files
Output: s3 bucket path to deposit output .docx files
"""
import json
import boto3
from botocore.exceptions import ClientError
import logging

from collections import deque
from functools import partial
from operator import itemgetter

from docx import Document
from docx.shared import Inches

def load_json(file_obj):
    data = json.loads(file_obj['Body'].read().decode('utf-8'))
    return data

def load_raw_transcription(file_obj):
    data = load_json(file_obj)
    status = data['status']

    if status != "COMPLETED":
        raise ValueError("AWS job [{}] status is {}".format(fname, status))

    return data['results'], data['jobName']

def get_end_times(speaker_label_items):
    items = sorted(speaker_label_items, key=lambda _: float(itemgetter('end_time')(_)))
    return deque([(i['end_time'], i['speaker_label']) for i in items])

def update_speaker(document, speaker_name, line_tstamp, speaker_map=()):
    speaker = speaker_map[int(speaker_name[-1])]
    tstamp = line_tstamp

    # Assign formatting for speaker and timestamp
    p = document.add_paragraph()
    run1 = p.add_run('{}:'.format(speaker))
    run1.bold = True
    run2 = p.add_run(' [{:.2f} mins]'.format(tstamp/60))
    run2.italic = True

    return document, speaker

def append(html, text, separator='\n', postfix=''):
    return '{}{}{}{}'.format(html, separator, text, postfix)

def get_token(line):
    if len(line['alternatives']) == 1:
        token = line['alternatives'][0]['content']
    else:
        token = sorted(line['alternatives'],
                       key=lambda _: float(itemgetter('confidence')(_)),
                       reverse=True)[0]
    return token

def get_speaker(end_time_speaker_dict, end_times_flt, line_start, line_end):
    end_times_sub = [i for i in end_times_flt if i >= line_end]
    end_times_sub.sort()
    closest_end = end_times_sub[0]
    speaker = end_time_speaker_dict[str(closest_end)]
    return speaker, closest_end

def build_worddoc(lines, end_times, job_name, speaker_map):
    document = Document()
    print(job_name)
    if '_' in job_name:
        job_name = job_name.replace("_", " ")

    document.add_heading('{}'.format(job_name), 0)
    _update_speaker = partial(update_speaker, speaker_map=speaker_map)
    current_speaker = None
    speaker_changes = 0

    end_time_speaker_dict = dict(end_times)
    end_times_flt = list(map(float, list(end_time_speaker_dict.keys())))
    line_tstamp = 0

    punctuation = ''
    doc_line = ''
    num_lines = len(lines)
    i = 0
    range_speaker = 'spk_0'
    
    while i < num_lines - 1:

        line = lines[i]
        # Get punctuation and append to doc_line; continue to next indexed line
        if line['type'] == 'punctuation':
            try:
                punctuation = append(punctuation, line['alternatives'][0]['content'], separator='', postfix=' ')
                doc_line += punctuation
                punctuation = ''
                i += 1
            except (KeyError, IndexError):
                pass
            continue

        #Check new line trigger:
        #Check for line-ending punctuation in last line's punctuation:
        if lines[i-1]['type'] == 'punctuation' and lines[i-1]['alternatives'][0]['content'] in '.!?':
            #Write string to document as paragraph:
            document, current_speaker = _update_speaker(document, range_speaker, line_tstamp)
            document.add_paragraph(doc_line)
            doc_line = ''

        #Otherwise get current line information:
        line_start = float(line['start_time'])
        line_end = float(line['end_time'])
        range_speaker, closest_end = get_speaker(end_time_speaker_dict,
                                                end_times_flt,
                                                line_start,
                                                line_end
                                                )
        #Get and assign token to doc_line string
        token = get_token(line)
        if len(doc_line) > 0:
            doc_line += ' ' + token
        else:
            line_tstamp = line_start
            doc_line = token

        #Update index:
        i += 1
        continue

    document.add_page_break()

    return document

def upload_file(file_name, bucket, object_name=None):
    """Upload a file to an S3 bucket

    :param file_name: File to upload
    :param bucket: Bucket to upload to
    :param object_name: S3 object name. If not specified then file_name is used
    :return: True if file was uploaded, else False
    """

    # If S3 object_name was not specified, use file_name
    if object_name is None:
        object_name = file_name

    # Upload the file
    s3_client = boto3.client('s3')
    try:
        response = s3_client.upload_file(file_name, bucket, object_name)
    except ClientError as e:
        logging.error(e)
        return False
    return True

def write_to_doc(document, fname, bucket_out_nm):
    filename = '{}.docx'.format(fname)
    document.save(filename)
    upload_file(file_name = filename, 
        bucket = bucket_out_nm)
    return

def parse_raw_transcription(fname, bucket_out_nm):
    data, job_name = load_raw_transcription(fname)
    n_speakers = data['speaker_labels']['speakers']
    speaker_names = ['speaker_' + str(i) for i in range(n_speakers)]
    end_times = get_end_times(data['speaker_labels']['segments'])
    lines = data['items']
    docx = build_worddoc(lines, end_times, job_name, speaker_names)
    write_to_doc(docx, job_name, bucket_out_nm)
    return


if __name__ == '__main__':
    s3 = boto3.resource('s3')
    s3_client = boto3.client('s3')
    bucket_in_nm = 'intermediate-transcribe-rwitt-research'
    prefix = 'medical'
    bucket_out_nm = 'output-transcript-rwitt-research'
    bucket_in = s3.Bucket(bucket_in_nm)
    
    # Get only .json files
    paginator = s3_client.get_paginator('list_objects_v2')
    response_iterator = paginator.paginate(Bucket=bucket_in_nm, Prefix=prefix)
    file_names = []

    for response in response_iterator:
        for object_data in response['Contents']:
            key = object_data['Key']
            if key.endswith('.json'):
                file_names.append(key)
    
    # Iterate through input s3 bucket 
    for obj in bucket_in.objects.all():
        fname = obj.key
        if fname in file_names:
            file_obj = s3_client.get_object(Bucket=bucket_in_nm, Key=fname)
            
            #Execute parsing code 
            parse_raw_transcription(file_obj, bucket_out_nm)
        
    print('Success!')